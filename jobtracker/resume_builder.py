"""Tailored resume builder using OpenAI."""

import os
from pathlib import Path
from typing import Optional, Tuple


def extract_text_from_file(filepath: str) -> Tuple[str, Optional[str]]:
    """
    Extract text from a .docx or .pdf file.
    Returns (extracted_text, error_message).
    If successful, error_message is None.
    """
    path = Path(filepath)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts).strip(), None
        except ImportError:
            return "", "Install PyMuPDF: pip install pymupdf"
        except Exception as e:
            return "", str(e)
    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(filepath)
            text_parts = [p.text for p in doc.paragraphs]
            return "\n".join(text_parts).strip(), None
        except ImportError:
            return "", "Install python-docx: pip install python-docx"
        except Exception as e:
            return "", str(e)
    else:
        return "", "Unsupported format. Use .pdf or .docx"

CONFIG_PATH = Path(__file__).parent.parent / "config.json"


def _load_config() -> dict:
    """Load config from JSON file."""
    import json
    if CONFIG_PATH.exists():
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            pass
    return {}


def _save_config(config: dict) -> None:
    """Save config to JSON file."""
    import json
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2)
    except OSError:
        pass


def get_api_key() -> Optional[str]:
    """Get OpenAI API key from config or environment."""
    config = _load_config()
    key = config.get("openai_api_key", "").strip()
    if key:
        return key
    return os.environ.get("OPENAI_API_KEY", "").strip() or None


def get_base_resume() -> str:
    """Get user's base resume from config."""
    return _load_config().get("base_resume", "")


def save_api_key(key: str) -> None:
    """Save OpenAI API key to config."""
    config = _load_config()
    config["openai_api_key"] = key
    _save_config(config)


def save_base_resume(resume: str) -> None:
    """Save base resume to config."""
    config = _load_config()
    config["base_resume"] = resume
    _save_config(config)


SECTION_NAMES = {
    "experience", "education", "skills", "summary", "objective",
    "certifications", "projects", "references", "work experience",
    "professional experience", "technical skills", "summary of qualifications",
    "employment", "employment history", "work history", "academic",
    "languages", "honors", "achievements", "activities", "volunteer",
    "publications", "interests", "additional information", "core competencies",
    "relevant coursework", "technical expertise", "key competencies",
    "research experience", "teaching experience", "leadership",
}


def parse_resume_sections(base_resume: str) -> list[str]:
    """
    Parse the base resume to extract section headers in order.
    Returns list of section names as they appear (e.g. ["EXPERIENCE", "EDUCATION", "SKILLS"]).
    Uses known section names; also captures custom ALL CAPS headers (excludes company-like names).
    """
    sections = []
    seen_lower = set()
    lines = base_resume.split("\n")
    # Skip company indicators - these are not section headers
    company_words = {"inc", "corp", "llc", "ltd", "co.", "company"}
    for line in lines:
        stripped = line.strip().lstrip("#*_- ").rstrip(":").strip()
        if not stripped or len(stripped) > 50:
            continue
        cleaned = stripped.lower().rstrip(":")
        # Known section
        if cleaned in SECTION_NAMES and cleaned not in seen_lower:
            sections.append(stripped)
            seen_lower.add(cleaned)
            continue
        # Custom section: ALL CAPS, 4-40 chars, no company indicators
        if (
            4 <= len(stripped) <= 40
            and stripped.replace(" ", "").isalpha()
            and stripped.isupper()
            and cleaned not in seen_lower
        ):
            words = set(cleaned.split())
            if not words & company_words:
                sections.append(stripped)
                seen_lower.add(cleaned)
    return sections


def parse_projects_section(base_resume: str) -> Optional[str]:
    """
    Extract the PROJECTS section content from the base resume.
    Returns the raw text of the projects section, or None if not found.
    """
    lines = base_resume.split("\n")
    in_projects = False
    project_lines = []
    project_headers = {"projects", "project experience", "personal projects", "selected projects"}

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_projects:
                project_lines.append("")  # Preserve blank lines within section
            continue
        cleaned = stripped.lower().lstrip("#*_- ").rstrip(":").strip()
        # Check if this is a section header
        if cleaned in project_headers or (cleaned.startswith("project") and len(cleaned) < 30):
            in_projects = True
            continue  # Skip the "PROJECTS" header line itself
        if in_projects:
            # Stop at next known section (not projects)
            if cleaned in SECTION_NAMES and cleaned not in project_headers:
                break
            if (cleaned in SECTION_NAMES or
                    (stripped.isupper() and 4 <= len(stripped) <= 45 and stripped.replace(" ", "").isalpha())):
                if cleaned not in project_headers:
                    break
            project_lines.append(stripped)
    return "\n".join(project_lines).strip() if project_lines else None


def build_tailored_resume(
    job_company: str,
    job_position: str,
    job_description: Optional[str],
    base_resume: str,
    api_key: Optional[str] = None,
) -> tuple[str, Optional[str]]:
    """
    Use OpenAI to generate a tailored resume for the given job.
    Returns (tailored_resume_text, error_message).
    If successful, error_message is None.
    """
    try:
        from openai import OpenAI
    except ImportError:
        return "", "Install openai package: pip install openai"

    key = api_key or get_api_key()
    if not key:
        return "", "OpenAI API key not set. Add it in Settings."

    if not base_resume.strip():
        return "", "Base resume not set. Add your resume in Settings."

    job_context = f"Company: {job_company}\nPosition: {job_position}"
    if job_description and job_description.strip():
        job_context += f"\n\nJob Description:\n{job_description.strip()}"

    # Parse projects from base resume (before sections, used in instruction)
    projects_content = parse_projects_section(base_resume)
    if projects_content:
        job_context += f"\n\nProjects from base resume (consider and connect these to the job):\n{projects_content}"

    # Parse base resume to get its section structure
    sections = parse_resume_sections(base_resume)
    section_instruction = ""
    if sections:
        section_list = ", ".join(sections)
        section_instruction = (
            f"IMPORTANT: The base resume has these sections in this exact order: {section_list}. "
            "Generate the tailored resume with EXACTLY these sections, in this order, using the same section header names. "
            "Do NOT add, remove, or reorder sections. "
        )

    project_instruction = ""
    if projects_content:
        project_instruction = (
            "The base resume includes a PROJECTS section. Consider these projects when tailoring: "
            "connect relevant projects to the job requirements, emphasize technologies and outcomes that match the role, "
            "and tailor project descriptions to align with the job. Do NOT omit the projects section. "
        )

    client = OpenAI(api_key=key)

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert resume writer. Tailor the given resume to match the job posting. "
                    "Emphasize relevant skills, experiences, and keywords from the job description. "
                    + section_instruction +
                    project_instruction +
                    "Do NOT use company or employer names as section headers - those go under experience entries. "
                    "Use this format:\n"
                    "- Section headers in ALL CAPS matching the base resume exactly\n"
                    "- For each job: Company name on its own line, then Role | Date on next line\n"
                    "- Date format: 'Nov 2024 - Dec 2025' or 'Nov 2024 - Present' for current roles (use month abbrev + year)\n"
                    "- Bullet points with • or - for achievements\n"
                    "- Blank line between different jobs (career breaks)\n"
                    "Output only the tailored resume text, no explanations.",
                },
                {
                    "role": "user",
                    "content": f"Job details:\n{job_context}\n\n---\n\nMy current resume:\n{base_resume}",
                },
            ],
        )
        content = response.choices[0].message.content
        return (content or "").strip(), None
    except Exception as e:
        return "", str(e)


def export_to_docx(text: str, filepath: str) -> Tuple[bool, Optional[str]]:
    """
    Export resume text to a Word (.docx) file with proper formatting.
    Detects section headers, bullet points, and applies appropriate styles.
    Name and contact info at the top are centered.
    """
    try:
        import re
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def _add_divider(paragraph):
            """Add a bottom border (horizontal line) to a paragraph."""
            p_pr = paragraph._element.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "000000")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)

        doc = Document()
        section_names = SECTION_NAMES

        def _clean_for_header(s: str) -> str:
            return s.strip().lstrip("#*_- ").rstrip(":").strip().lower()

        def _is_section_header(stripped: str, cleaned: str) -> bool:
            # Only known section names - never treat company names as sections
            if not cleaned:
                return False
            base = cleaned.rstrip(":")
            if base in section_names:
                return True
            if stripped.startswith("##"):
                after_hash = _clean_for_header(stripped[2:])
                return after_hash.rstrip(":") in section_names
            return False

        lines = text.split("\n")
        # Find where header block (name + contact) ends - first section header or bullet
        header_end = 0
        for idx, line in enumerate(lines):
            s = line.strip()
            if not s:
                continue
            cleaned = _clean_for_header(s)
            bullet_chars = ("•", "-", "*", "·", "▪", "▸")
            is_bullet = any(s.startswith(c) for c in bullet_chars) or (len(s) > 2 and s[0].isdigit() and s[1] in ".)")
            if _is_section_header(s, cleaned) or is_bullet:
                header_end = idx
                break
            header_end = idx + 1

        # Output header block (name + contact) centered
        header_lines = [l.strip() for l in lines[:header_end] if l.strip()]
        for idx, stripped in enumerate(header_lines):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            if idx == 0:
                run.bold = True
                run.font.size = Pt(16)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_lines:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Process rest of document
        prev_was_bullet = False
        month_abbrevs = r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)"

        def _normalize_date_range(text: str) -> str:
            """Fix date ranges: 'Mar 2022 Aug 2024' -> 'Mar 2022 - Aug 2024'."""
            # Two month-year pairs with only space(s) between (no hyphen)
            p_month_year = re.compile(
                rf"({month_abbrevs}\s+(?:19|20)\d{{2}})\s+({month_abbrevs}\s+(?:19|20)\d{{2}})",
                re.I
            )
            # Month-year followed by Present
            p_present = re.compile(
                rf"({month_abbrevs}\s+(?:19|20)\d{{2}})\s+(present)",
                re.I
            )
            result = p_month_year.sub(r"\1 - \2", text)
            result = p_present.sub(r"\1 - \2", result)
            return result

        def _parse_role_date_line(s: str) -> tuple[bool, str, str, str]:
            """Parse 'Role | Mar 2022 - Aug 2024'. Position=bold italic, date=normal only."""
            # 1) Try explicit separators (| or ,) - role and date clearly split
            for sep in (" | ", "|", " – ", "–", ", "):
                if sep in s:
                    parts = s.split(sep, 1)
                    if len(parts) == 2:
                        role_part, date_part = parts[0].strip(), parts[1].strip()
                        if not role_part or not date_part:
                            continue
                        if not (re.search(r"(?:19|20)\d{2}|present", date_part, re.I) or
                                re.search(month_abbrevs, date_part, re.I)):
                            continue
                        if re.search(r"(?:19|20)\d{2}\s*$", role_part):
                            continue
                        date_part = _normalize_date_range(date_part)
                        disp = " | " if sep in ("|", " | ") else (" – " if sep in ("–", " – ") else ", ")
                        return True, role_part, date_part, disp
            # 2) No separator: "Software Engineer Mar 2022 - Aug 2024" - find where date starts
            m = re.search(rf"({month_abbrevs}\s+(?:19|20)\d{{2}}\s*[-–—]?\s*(?:{month_abbrevs}\s+(?:19|20)\d{{2}}|present))", s, re.I)
            if m:
                date_part = _normalize_date_range(m.group(1).strip())
                role_part = s[:m.start()].strip().rstrip(",-–— ")
                if role_part and len(role_part) > 2:
                    return True, role_part, date_part, " | "
            return False, s, "", ""

        i = header_end
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if not stripped:
                i += 1
                continue

            cleaned = _clean_for_header(stripped)
            is_header = _is_section_header(stripped, cleaned)
            bullet_chars = ("•", "-", "*", "·", "▪", "▸")
            is_bullet = (
                any(stripped.startswith(c) for c in bullet_chars)
                or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)")
            )
            if is_bullet:
                bullet_text = stripped.lstrip("•-*·▪▸ ").lstrip("0123456789.) ")
                if not bullet_text and stripped:
                    bullet_text = stripped
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(bullet_text)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Pt(18)
                prev_was_bullet = True
            elif is_header:
                p = doc.add_paragraph()
                header_text = stripped.lstrip("#*_- ").rstrip(":").strip()
                run = p.add_run(header_text)
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                _add_divider(p)
                prev_was_bullet = False
            else:
                # Company name or role line
                is_role, role_part, date_part, sep = _parse_role_date_line(stripped)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                if prev_was_bullet:
                    p.paragraph_format.space_before = Pt(12)  # Paragraph between career breaks
                if is_role and role_part and date_part:
                    run_role = p.add_run(role_part)
                    run_role.bold = True
                    run_role.italic = True
                    run_date = p.add_run(sep + date_part)
                    run_date.bold = False
                    run_date.italic = False
                else:
                    run = p.add_run(stripped)
                    if re.search(rf"{month_abbrevs}\s+(?:19|20)\d{{2}}|(?:19|20)\d{{2}}\s*[-–—]", stripped, re.I):
                        run.bold = False  # Date line - never bold
                    else:
                        run.bold = True  # Company name
                prev_was_bullet = False

            i += 1

        # Add normal style for body if document has no styles
        doc.save(filepath)
        return True, None
    except ImportError:
        return False, "Install python-docx: pip install python-docx"
    except Exception as e:
        return False, str(e)
